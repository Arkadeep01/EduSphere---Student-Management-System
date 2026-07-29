import json
from io import BytesIO

from django.db import transaction
from django.utils import timezone
from django.http import HttpResponse, FileResponse

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import status

from administration.permissions import IsAdmin, IsAdminOrDirector, IsDirector
from administration.models.template import DocumentTemplate, GeneratedDocument
from administration.serializers.template import (
    DocumentTemplateListSerializer,
    DocumentTemplateDetailSerializer,
    DocumentTemplateCreateSerializer,
    DocumentTemplateActivateSerializer,
    GeneratedDocumentListSerializer,
    GeneratedDocumentDetailSerializer,
    GenerateDocumentSerializer,
    PlaceholderValidateSerializer,
)
from administration.services.template_service import DocumentTemplateService, MyDocumentService, PLACEHOLDER_REGISTRY, PLACEHOLDER_PATTERN
from authentication.models import CustomUser
from student.models import StudentProfile


# ── Document Templates ─────────────────────────────────────────────


class TemplateListView(APIView):
    permission_classes = [IsAuthenticated, IsAdminOrDirector]

    def get(self, request):
        templates = DocumentTemplate.objects.all().select_related("created_by", "approved_by")
        serializer = DocumentTemplateListSerializer(templates, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = DocumentTemplateCreateSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        template = DocumentTemplate.objects.create(
            name=serializer.validated_data["name"],
            document_type=serializer.validated_data["document_type"],
            description=serializer.validated_data.get("description", ""),
            created_by=request.user,
        )
        result = DocumentTemplateDetailSerializer(template)
        return Response(result.data, status=status.HTTP_201_CREATED)


class TemplateDetailView(APIView):
    permission_classes = [IsAuthenticated, IsAdminOrDirector]

    def get(self, request, template_id):
        try:
            template = DocumentTemplate.objects.select_related("created_by", "approved_by").get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)
        serializer = DocumentTemplateDetailSerializer(template)
        return Response(serializer.data)

    def patch(self, request, template_id):
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)
        if template.status == "active":
            return Response({"error": "Active templates cannot be edited. Retire then create a new version."},
                            status=status.HTTP_400_BAD_REQUEST)
        if template.status == "retired":
            return Response({"error": "Retired templates cannot be edited."},
                            status=status.HTTP_400_BAD_REQUEST)
        for field in ["name", "description"]:
            if field in request.data:
                setattr(template, field, request.data[field])
        template.save()
        serializer = DocumentTemplateDetailSerializer(template)
        return Response(serializer.data)

    def delete(self, request, template_id):
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)
        if GeneratedDocument.objects.filter(template=template).exists():
            return Response({"error": "Cannot delete template with generated documents. Retire it instead."},
                            status=status.HTTP_400_BAD_REQUEST)
        template.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class TemplateUploadFileView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def post(self, request, template_id):
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)

        if template.status not in ("draft", "pending_approval"):
            return Response({"error": "Only draft or pending-approval templates can be updated."},
                            status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES.get("file")
        if not file:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)
        if not file.name.lower().endswith(".docx"):
            return Response({"error": "Only DOCX files are supported."}, status=status.HTTP_400_BAD_REQUEST)

        template.file = file

        try:
            content = self._read_docx_text(file)
        except Exception as e:
            return Response({"error": f"Failed to read template: {e}"}, status=status.HTTP_400_BAD_REQUEST)

        found_placeholders = DocumentTemplateService.extract_placeholders(content)
        validation = DocumentTemplateService.validate_placeholders(found_placeholders)

        template.placeholder_registry = validation["valid"]
        if validation["unknown"]:
            template.status = "draft"
            template.save()
            return Response({
                "error": "Template contains unknown placeholders",
                "unknown_placeholders": validation["unknown"],
                "valid_placeholders": validation["valid"],
                "placeholders": found_placeholders,
            }, status=status.HTTP_400_BAD_REQUEST)

        if template.status == "draft":
            template.status = "pending_approval"

        file.seek(0)
        template.save()
        serializer = DocumentTemplateDetailSerializer(template)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @staticmethod
    def _read_docx_text(file):
        from docx import Document
        doc = Document(file)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)
        try:
            for section in doc.sections:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
        except Exception:
            pass
        return "\n".join(parts)


class TemplateActivateView(APIView):
    permission_classes = [IsAuthenticated, IsDirector]

    def post(self, request, template_id):
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = DocumentTemplateActivateSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        action = serializer.validated_data["action"]
        reason = serializer.validated_data.get("reason", "")

        if action == "submit":
            if template.status != "draft":
                return Response({"error": "Only draft templates can be submitted."}, status=status.HTTP_400_BAD_REQUEST)
            if not template.file:
                return Response({"error": "Upload template file before submitting."}, status=status.HTTP_400_BAD_REQUEST)
            template.status = "pending_approval"
            template.save()
            return Response({"message": "Template submitted for approval."})

        if action == "approve":
            if template.status != "pending_approval":
                return Response({"error": "Only pending-approval templates can be approved."},
                                status=status.HTTP_400_BAD_REQUEST)
            if not template.file:
                return Response({"error": "Template file is missing."}, status=status.HTTP_400_BAD_REQUEST)

            old_active = DocumentTemplate.objects.filter(
                document_type=template.document_type, status="active"
            ).first()
            if old_active and old_active.id != template.id:
                old_active.status = "retired"
                old_active.save()

            template.status = "active"
            template.approved_by = request.user
            template.approved_at = timezone.now()
            template.rejection_reason = ""
            template.save()
            return Response({"message": "Template approved and activated."})

        if action == "reject":
            if template.status != "pending_approval":
                return Response({"error": "Only pending-approval templates can be rejected."},
                                status=status.HTTP_400_BAD_REQUEST)
            template.status = "draft"
            template.rejection_reason = reason
            template.save()
            return Response({"message": "Template rejected."})

        if action == "retire":
            if template.status != "active":
                return Response({"error": "Only active templates can be retired."},
                                status=status.HTTP_400_BAD_REQUEST)
            template.status = "retired"
            template.save()
            return Response({"message": "Template retired."})

        return Response({"error": "Invalid action."}, status=status.HTTP_400_BAD_REQUEST)


class TemplatePlaceholdersView(APIView):
    permission_classes = [IsAuthenticated, IsAdminOrDirector]

    def get(self, request):
        registry = [
            {"placeholder": k, "label": v["label"], "domain": v["domain"], "type": v["type"]}
            for k, v in sorted(PLACEHOLDER_REGISTRY.items())
        ]
        return Response(registry)

    def post(self, request, template_id):
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)
        if not template.file:
            return Response({"error": "Upload template file first."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            content = TemplateUploadFileView._read_docx_text(template.file)
        except Exception as e:
            return Response({"error": f"Failed to read template: {e}"}, status=status.HTTP_400_BAD_REQUEST)

        found = DocumentTemplateService.extract_placeholders(content)
        validation = DocumentTemplateService.validate_placeholders(found)

        return Response({
            "found_placeholders": found,
            "found_count": len(found),
            "valid_placeholders": validation["valid"],
            "unknown_placeholders": validation["unknown"],
            "is_valid": validation["is_valid"],
        })


class TemplatePreviewView(APIView):
    permission_classes = [IsAuthenticated, IsAdminOrDirector]

    def post(self, request, template_id):
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found"}, status=status.HTTP_404_NOT_FOUND)
        if not template.file:
            return Response({"error": "Template file not uploaded."}, status=status.HTTP_400_BAD_REQUEST)

        context_data = request.data.get("context_data", {})
        full_context = DocumentTemplateService.resolve_data_context(context_data)

        try:
            docx_bytes = DocumentTemplateService.process_docx_template(template.file, full_context)
        except Exception as e:
            return Response({"error": f"Template processing failed: {e}"}, status=status.HTTP_400_BAD_REQUEST)

        return HttpResponse(docx_bytes, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            headers={"Content-Disposition": f"inline; filename={template.name}_preview.docx"})


# ── Generated Documents ────────────────────────────────────────────


class GeneratedDocumentListView(APIView):
    permission_classes = [IsAuthenticated, IsAdminOrDirector]

    def get(self, request):
        docs = GeneratedDocument.objects.all().select_related("template", "generated_by")
        document_type = request.query_params.get("document_type")
        if document_type:
            docs = docs.filter(document_type=document_type)
        serializer = GeneratedDocumentListSerializer(docs, many=True)
        return Response(serializer.data)


class GeneratedDocumentCreateView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def post(self, request):
        serializer = GenerateDocumentSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        data = serializer.validated_data
        recipient_user = None
        if data.get("recipient_user_id"):
            try:
                recipient_user = CustomUser.objects.get(id=data["recipient_user_id"])
            except CustomUser.DoesNotExist:
                return Response({"error": "Recipient user not found"}, status=status.HTTP_404_NOT_FOUND)

        try:
            doc = DocumentTemplateService.generate_document(
                template_id=data["template_id"],
                context_data=data.get("context_data", {}),
                generated_by=request.user,
                recipient_user=recipient_user,
                recipient_name=data.get("recipient_name", ""),
                recipient_entity=data.get("recipient_entity", ""),
                academic_session=data.get("academic_session", ""),
                output_format=data.get("output_format", "docx"),
            )
        except DocumentTemplate.DoesNotExist:
            return Response({"error": "Template not found or not active"}, status=status.HTTP_404_NOT_FOUND)
        except ValueError as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

        result = GeneratedDocumentDetailSerializer(doc)
        return Response(result.data, status=status.HTTP_201_CREATED)


class GeneratedDocumentDetailView(APIView):
    permission_classes = [IsAuthenticated, IsAdminOrDirector]

    def get(self, request, doc_id):
        try:
            doc = GeneratedDocument.objects.select_related("template", "generated_by").get(id=doc_id)
        except GeneratedDocument.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        serializer = GeneratedDocumentDetailSerializer(doc)
        return Response(serializer.data)

    def delete(self, request, doc_id):
        if request.user.role not in ("admin", "director"):
            return Response({"error": "Access denied"}, status=status.HTTP_403_FORBIDDEN)
        try:
            doc = GeneratedDocument.objects.get(id=doc_id)
        except GeneratedDocument.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        doc.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class GeneratedDocumentDownloadView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, doc_id):
        try:
            doc = GeneratedDocument.objects.get(id=doc_id)
        except GeneratedDocument.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)

        if doc.recipient_user and doc.recipient_user_id != request.user.id and request.user.role not in ("admin", "director"):
            return Response({"error": "Access denied"}, status=status.HTTP_403_FORBIDDEN)

        if not doc.file:
            return Response({"error": "File not found"}, status=status.HTTP_404_NOT_FOUND)

        content_type = "application/pdf" if doc.file_format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(doc.file, content_type=content_type,
                            headers={"Content-Disposition": f"attachment; filename={doc.reference_number}.{doc.file_format}"})


# ── My Documents (Recipient facing) ────────────────────────────────


class MyDocumentListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        docs = MyDocumentService.list_received_documents(request.user)
        serializer = GeneratedDocumentListSerializer(docs, many=True)
        return Response(serializer.data)
