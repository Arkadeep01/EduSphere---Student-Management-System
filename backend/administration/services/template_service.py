import re
import os
import uuid
from io import BytesIO
from datetime import datetime

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from django.template import Template, Context
from django.utils import timezone

from administration.models.template import DocumentTemplate, GeneratedDocument
from administration.models.academic import AcademicSession
from student.models import StudentProfile, Subject
from teacher.models import TeacherProfile
from administration.models.letterhead import Letterhead


PLACEHOLDER_REGISTRY = {
    # Institution
    "institution.name": {"label": "Institution Name", "domain": "institution", "type": "string"},
    "institution.address": {"label": "Institution Address", "domain": "institution", "type": "string"},
    "institution.phone": {"label": "Institution Phone", "domain": "institution", "type": "string"},
    "institution.email": {"label": "Institution Email", "domain": "institution", "type": "string"},
    "institution.website": {"label": "Institution Website", "domain": "institution", "type": "string"},
    "institution.logo_url": {"label": "Institution Logo URL", "domain": "institution", "type": "url"},
    "institution.principal_name": {"label": "Principal Name", "domain": "institution", "type": "string"},
    # Academic Session
    "session.name": {"label": "Session Name", "domain": "session", "type": "string"},
    "session.start_date": {"label": "Session Start Date", "domain": "session", "type": "date"},
    "session.end_date": {"label": "Session End Date", "domain": "session", "type": "date"},
    # Student
    "student.name": {"label": "Student Full Name", "domain": "student", "type": "string"},
    "student.first_name": {"label": "Student First Name", "domain": "student", "type": "string"},
    "student.last_name": {"label": "Student Last Name", "domain": "student", "type": "string"},
    "student.roll_number": {"label": "Roll Number", "domain": "student", "type": "string"},
    "student.admission_number": {"label": "Admission Number", "domain": "student", "type": "string"},
    "student.class_assigned": {"label": "Class", "domain": "student", "type": "string"},
    "student.section": {"label": "Section", "domain": "student", "type": "string"},
    "student.father_name": {"label": "Father's Name", "domain": "student", "type": "string"},
    "student.mother_name": {"label": "Mother's Name", "domain": "student", "type": "string"},
    "student.guardian_contact": {"label": "Guardian Contact", "domain": "student", "type": "string"},
    "student.date_of_birth": {"label": "Date of Birth", "domain": "student", "type": "date"},
    "student.gender": {"label": "Gender", "domain": "student", "type": "string"},
    "student.email": {"label": "Student Email", "domain": "student", "type": "string"},
    "student.phone": {"label": "Student Phone", "domain": "student", "type": "string"},
    "student.address": {"label": "Student Address", "domain": "student", "type": "string"},
    # Academic Performance
    "result.exam_name": {"label": "Exam Name", "domain": "result", "type": "string"},
    "result.academic_year": {"label": "Academic Year", "domain": "result", "type": "string"},
    "result.total_marks_obtained": {"label": "Total Marks Obtained", "domain": "result", "type": "number"},
    "result.total_marks_max": {"label": "Total Marks Maximum", "domain": "result", "type": "number"},
    "result.percentage": {"label": "Percentage", "domain": "result", "type": "number"},
    "result.grade": {"label": "Overall Grade", "domain": "result", "type": "string"},
    "result.grade_point": {"label": "Grade Point", "domain": "result", "type": "number"},
    "result.merit_rank": {"label": "Merit Rank", "domain": "result", "type": "number"},
    "result.class_rank": {"label": "Class Rank", "domain": "result", "type": "number"},
    "result.remarks": {"label": "Remarks", "domain": "result", "type": "string"},
    "result.is_pass": {"label": "Pass Status", "domain": "result", "type": "string"},
    # Attendance
    "attendance.total_days": {"label": "Total Working Days", "domain": "attendance", "type": "number"},
    "attendance.present_days": {"label": "Days Present", "domain": "attendance", "type": "number"},
    "attendance.absent_days": {"label": "Days Absent", "domain": "attendance", "type": "number"},
    "attendance.percentage": {"label": "Attendance Percentage", "domain": "attendance", "type": "number"},
    # Fee
    "fee.total_fee": {"label": "Total Fee", "domain": "fee", "type": "number"},
    "fee.paid_amount": {"label": "Amount Paid", "domain": "fee", "type": "number"},
    "fee.pending_amount": {"label": "Pending Amount", "domain": "fee", "type": "number"},
    "fee.receipt_number": {"label": "Receipt Number", "domain": "fee", "type": "string"},
    "fee.payment_date": {"label": "Payment Date", "domain": "fee", "type": "date"},
    "fee.academic_session": {"label": "Fee Session", "domain": "fee", "type": "string"},
    # Teacher
    "teacher.name": {"label": "Teacher Full Name", "domain": "teacher", "type": "string"},
    "teacher.employee_id": {"label": "Employee ID", "domain": "teacher", "type": "string"},
    "teacher.email": {"label": "Teacher Email", "domain": "teacher", "type": "string"},
    "teacher.phone": {"label": "Teacher Phone", "domain": "teacher", "type": "string"},
    "teacher.qualification": {"label": "Qualification", "domain": "teacher", "type": "string"},
    "teacher.designation": {"label": "Designation", "domain": "teacher", "type": "string"},
    "teacher.assigned_subject": {"label": "Assigned Subject", "domain": "teacher", "type": "string"},
    # Signature & Date
    "signature.principal": {"label": "Principal Signature", "domain": "system", "type": "string"},
    "signature.principal_name": {"label": "Principal Name (Printed)", "domain": "system", "type": "string"},
    "signature.date": {"label": "Current Date", "domain": "system", "type": "date"},
    "signature.date_long": {"label": "Current Date (Long Format)", "domain": "system", "type": "date"},
    "signature.place": {"label": "Place (City)", "domain": "system", "type": "string"},
    # General
    "general.current_date": {"label": "Current Date", "domain": "system", "type": "date"},
    "general.current_date_long": {"label": "Current Date (Long)", "domain": "system", "type": "date"},
    "general.reference_number": {"label": "Document Reference Number", "domain": "system", "type": "string"},
}

VALID_PLACEHOLDER_NAMES = set(PLACEHOLDER_REGISTRY.keys())

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_.]*)\s*\}\}")


class DocumentTemplateService:

    @staticmethod
    def extract_placeholders(content: str) -> list:
        found = set(PLACEHOLDER_PATTERN.findall(content))
        return sorted(found)

    @staticmethod
    def validate_placeholders(placeholders: list) -> dict:
        unknown = [p for p in placeholders if p not in VALID_PLACEHOLDER_NAMES]
        valid = [p for p in placeholders if p in VALID_PLACEHOLDER_NAMES]
        return {
            "valid": valid,
            "valid_count": len(valid),
            "unknown": unknown,
            "unknown_count": len(unknown),
            "is_valid": len(unknown) == 0,
        }

    @staticmethod
    def resolve_data_context(context_data: dict) -> dict:
        ctx = {}
        # Institution
        letterhead = Letterhead.objects.filter(is_default=True).first() or Letterhead.objects.first()
        if letterhead:
            branding = letterhead.branding if isinstance(letterhead.branding, dict) else {}
            ctx["institution.name"] = branding.get("school_name", "")
            ctx["institution.address"] = branding.get("address", "")
            ctx["institution.phone"] = branding.get("phone", "")
            ctx["institution.email"] = branding.get("email", "")
            ctx["institution.website"] = branding.get("website", "")
            ctx["institution.logo_url"] = (
                f"{settings.MEDIA_URL}{letterhead.logo.name}" if letterhead.logo else ""
            )
            ctx["institution.principal_name"] = branding.get("principal_name", "")
            ctx["signature.principal_name"] = branding.get("principal_name", "")
            ctx["signature.place"] = branding.get("city", "")

        # Session
        session = context_data.get("academic_session", "")
        if session:
            ctx["session.name"] = session
        else:
            active_session = AcademicSession.objects.filter(is_current=True).first()
            ctx["session.name"] = str(active_session) if active_session else ""

        # Student
        student_id = context_data.get("student_id")
        if student_id:
            try:
                student = StudentProfile.objects.select_related("user").get(id=student_id)
                user = student.user
                ctx["student.name"] = user.get_full_name() or user.email
                ctx["student.first_name"] = user.first_name or ""
                ctx["student.last_name"] = user.last_name or ""
                ctx["student.roll_number"] = student.roll_number or ""
                ctx["student.admission_number"] = student.admission_number or ""
                ctx["student.class_assigned"] = student.class_assigned or ""
                ctx["student.section"] = student.section or ""
                ctx["student.father_name"] = student.father_name or ""
                ctx["student.mother_name"] = student.mother_name or ""
                ctx["student.guardian_contact"] = student.guardian_contact or ""
                ctx["student.date_of_birth"] = str(student.date_of_birth) if student.date_of_birth else ""
                ctx["student.gender"] = student.gender or ""
                ctx["student.email"] = user.email or ""
                ctx["student.phone"] = student.phone or ""
                ctx["student.address"] = student.address or ""
            except StudentProfile.DoesNotExist:
                pass

        # Teacher
        teacher_id = context_data.get("teacher_id")
        if teacher_id:
            try:
                teacher = TeacherProfile.objects.select_related("user").get(id=teacher_id)
                ctx["teacher.name"] = teacher.user.get_full_name() or teacher.user.email
                ctx["teacher.employee_id"] = teacher.employee_id or ""
                ctx["teacher.email"] = teacher.user.email or ""
                ctx["teacher.phone"] = teacher.phone or ""
                ctx["teacher.qualification"] = teacher.qualification or ""
                ctx["teacher.designation"] = teacher.designation or ""
            except TeacherProfile.DoesNotExist:
                pass

        # Result data
        for key in ["exam_name", "academic_year", "total_marks_obtained", "total_marks_max",
                     "percentage", "grade", "grade_point", "merit_rank", "class_rank",
                     "remarks", "is_pass"]:
            val = context_data.get(key)
            if val is not None:
                ctx[f"result.{key}"] = val

        # Attendance data
        for key in ["total_days", "present_days", "absent_days"]:
            val = context_data.get(key)
            if val is not None:
                ctx[f"attendance.{key}"] = val
        att_pct = context_data.get("attendance_percentage")
        if att_pct is not None:
            ctx["attendance.percentage"] = att_pct

        # Fee data
        for key in ["total_fee", "paid_amount", "pending_amount", "receipt_number",
                     "payment_date", "academic_session"]:
            val = context_data.get(key)
            if val is not None:
                ctx[f"fee.{key}"] = val

        # Signatures & Dates
        now = timezone.now()
        ctx["signature.date"] = now.strftime("%d-%m-%Y")
        ctx["signature.date_long"] = now.strftime("%d %B %Y")
        ctx["general.current_date"] = now.strftime("%d-%m-%Y")
        ctx["general.current_date_long"] = now.strftime("%d %B %Y")
        ctx["general.reference_number"] = context_data.get("reference_number", "")

        # Merge any user-provided overrides
        overrides = context_data.get("placeholders", {})
        if isinstance(overrides, dict):
            ctx.update(overrides)

        return ctx

    @staticmethod
    def generate_reference_number(document_type: str, session: str, scope: str, sequence: int) -> str:
        prefix = DocumentTemplate.DOC_TYPE_PREFIX_MAP.get(document_type, "CD")
        seq_str = f"{sequence:04d}"
        return f"{prefix}/{session}/{scope}/{seq_str}"

    @staticmethod
    def process_docx_template(template_file, context: dict) -> bytes:
        from docx import Document
        doc = Document(template_file)

        def replace_in_paragraph(paragraph):
            for run in paragraph.runs:
                text = run.text
                matches = PLACEHOLDER_PATTERN.findall(text)
                for placeholder in matches:
                    value = str(context.get(placeholder, f"{{{{{placeholder}}}}}"))
                    text = text.replace("{{" + placeholder + "}}", value)
                run.text = text

        def replace_in_cell(cell):
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph)

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_cell(cell)

        try:
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    replace_in_paragraph(paragraph)
                footer = section.footer
                for paragraph in footer.paragraphs:
                    replace_in_paragraph(paragraph)
        except Exception:
            pass

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    @staticmethod
    def get_next_sequence() -> int:
        last = GeneratedDocument.objects.order_by("id").last()
        if last:
            try:
                parts = last.reference_number.split("/")
                return int(parts[-1]) + 1
            except (IndexError, ValueError):
                pass
        return 1

    @classmethod
    def generate_document(cls, template_id, context_data, generated_by,
                          recipient_user=None, recipient_name="",
                          recipient_entity="", academic_session="",
                          output_format="docx"):
        template = DocumentTemplate.objects.get(id=template_id, status="active")
        sequence = cls.get_next_sequence()
        session_str = academic_session or str(AcademicSession.objects.filter(is_active=True).first() or "NA")
        scope = recipient_name or recipient_entity or str(recipient_user.id if recipient_user else "NA")
        ref_number = cls.generate_reference_number(
            template.document_type, session_str, scope, sequence
        )
        context_data["reference_number"] = ref_number
        full_context = cls.resolve_data_context(context_data)

        if not template.file:
            raise ValueError("Template file not found")

        docx_bytes = cls.process_docx_template(template.file, full_context)

        filename = f"{ref_number}.{output_format}"
        if output_format == "pdf":
            pdf_bytes = cls._convert_docx_to_pdf(docx_bytes, full_context)
            file_content = ContentFile(pdf_bytes, name=filename)
        else:
            file_content = ContentFile(docx_bytes, name=filename)

        doc = GeneratedDocument.objects.create(
            template=template,
            template_version=template.version,
            document_type=template.document_type,
            reference_number=ref_number,
            recipient_user=recipient_user,
            recipient_name=recipient_name,
            recipient_entity=recipient_entity,
            academic_session=academic_session,
            context_data=context_data,
            file=file_content,
            file_format=output_format,
            generated_by=generated_by,
        )
        return doc

    @staticmethod
    def _convert_docx_to_pdf(docx_bytes: bytes, context: dict) -> bytes:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_CENTER

            buf = BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("Title2", parent=styles["Title"], spaceAfter=12, alignment=TA_CENTER)
            normal_style = ParagraphStyle("Normal2", parent=styles["Normal"], spaceAfter=6, fontSize=11)

            from docx import Document as DocxDocument
            source = DocxDocument(BytesIO(docx_bytes))

            story = []
            for para in source.paragraphs:
                text = para.text.strip()
                if not text:
                    story.append(Spacer(1, 6))
                    continue
                style = title_style if para.style.name.startswith("Heading") else normal_style
                safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
                story.append(Paragraph(safe, style))

            for table in source.tables:
                story.append(Spacer(1, 8))
                table_data = []
                for row in table.rows:
                    table_data.append([Paragraph(cell.text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), normal_style) for cell in row.cells])
                from reportlab.platypus import Table as RLTable, TableStyle
                from reportlab.lib import colors
                t = RLTable(table_data)
                t.setStyle(TableStyle([
                    ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
                    ("BACKGROUND", (0,0), (-1,0), colors.Color(0.9, 0.9, 0.9)),
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ("LEFTPADDING", (0,0), (-1,-1), 6),
                    ("RIGHTPADDING", (0,0), (-1,-1), 6),
                    ("TOPPADDING", (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                ]))
                story.append(t)

            doc.build(story)
            return buf.getvalue()
        except ImportError as e:
            raise ValueError(f"PDF library not available: {e}")
        except Exception as e:
            raise ValueError(f"PDF conversion failed: {e}")


class MyDocumentService:
    @staticmethod
    def list_received_documents(user):
        return GeneratedDocument.objects.filter(
            recipient_user=user, is_archived=False
        ).select_related("template")
